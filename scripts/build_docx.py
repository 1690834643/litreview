#!/usr/bin/env python3
"""
build_docx.py — 把 Markdown 综述转成易读、用期刊标准字体的 .docx。

字体（Word 打开时本就内置，渲染无忧）：
  正文：西文 Times New Roman 12pt + 中文 宋体；行距 1.5；两端对齐
  标题：西文 Times New Roman 加粗 + 中文 黑体（H1 16 / H2 14 / H3 13）
  表格：西文 Times New Roman + 中文 宋体 10.5pt（五号，文献表更紧凑）

做法：用 python-docx 生成一个定义好上述样式的 reference 模板，再用 pandoc
套模板把 md 转 docx（pandoc 解析标题/列表/表格最稳）。无 pandoc 时退化为
python-docx 基础转换。

用法：
  python build_docx.py review.md -o review.docx
  python build_docx.py review.md -o review.docx --title "二化螟性别决定研究综述"
"""
import argparse
import os
import shutil
import subprocess
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

LATIN = "Times New Roman"
CJK_BODY = "宋体"
CJK_HEAD = "黑体"

ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets")
REF_DOC = os.path.normpath(os.path.join(ASSETS, "journal_reference.docx"))


def _set_style_font(style, latin, cjk, size_pt, bold=False, color=None):
    f = style.font
    f.name = latin
    f.size = Pt(size_pt)
    f.bold = bold
    if color:
        f.color.rgb = RGBColor(*color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def build_reference_docx(path=REF_DOC):
    """生成期刊标准字体的样式模板（pandoc --reference-doc 用）。"""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = Document()

    normal = doc.styles["Normal"]
    _set_style_font(normal, LATIN, CJK_BODY, 12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heads = {"Title": (18, CJK_HEAD), "Heading 1": (16, CJK_HEAD),
             "Heading 2": (14, CJK_HEAD), "Heading 3": (13, CJK_HEAD)}
    for name, (sz, cjk) in heads.items():
        if name in doc.styles:
            st = doc.styles[name]
            _set_style_font(st, LATIN, cjk, sz, bold=True,
                            color=(0x1F, 0x3A, 0x5F))
            st.paragraph_format.space_before = Pt(10)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing = 1.3

    # pandoc 会给单元格/首段/列表套这些样式名；全新 Document 里它们不存在，
    # 必须主动创建并设字体，否则会回退到主题字体（Calibri/Cambria）而非期刊字体。
    def ensure_para(name, size):
        st = doc.styles[name] if name in doc.styles \
            else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(st, LATIN, CJK_BODY, size)

    ensure_para("Body Text", 12)        # pandoc 正文段
    ensure_para("First Paragraph", 12)  # 标题后首段
    ensure_para("Compact", 10.5)        # 表格单元格 / 紧凑列表（五号）
    ensure_para("Table Caption", 10.5)
    ensure_para("Image Caption", 10.5)

    doc.save(path)
    return path


def via_pandoc(md_path, out_path, title=None):
    if not os.path.exists(REF_DOC):
        build_reference_docx(REF_DOC)
    cmd = ["pandoc", md_path, "-o", out_path,
           "--reference-doc", REF_DOC, "--toc", "--toc-depth=2"]
    if title:
        cmd += ["--metadata", f"title={title}"]
    subprocess.run(cmd, check=True)


def via_pythondocx(md_path, out_path, title=None):
    """无 pandoc 的兜底：极简 md→docx（标题/段落/无序列表）。"""
    doc = Document()
    _set_style_font(doc.styles["Normal"], LATIN, CJK_BODY, 12)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    if title:
        doc.add_heading(title, 0)
    for raw in open(md_path, encoding="utf-8"):
        line = raw.rstrip("\n")
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], 3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser(description="Markdown 综述 → 期刊标准字体 .docx")
    ap.add_argument("md", help="输入 Markdown 文件")
    ap.add_argument("-o", "--out", required=True, help="输出 .docx 路径")
    ap.add_argument("--title", help="文档标题（写入封面/元数据）")
    ap.add_argument("--rebuild-ref", action="store_true", help="强制重建字体模板")
    args = ap.parse_args()

    if args.rebuild_ref and os.path.exists(REF_DOC):
        os.remove(REF_DOC)

    if shutil.which("pandoc"):
        try:
            via_pandoc(args.md, args.out, args.title)
            print(f"# pandoc + 期刊字体模板 → {args.out}", file=sys.stderr)
            return
        except Exception as e:
            print(f"# pandoc 失败（{e}），退化 python-docx", file=sys.stderr)
    via_pythondocx(args.md, args.out, args.title)
    print(f"# python-docx 兜底 → {args.out}", file=sys.stderr)


if __name__ == "__main__":
    main()
